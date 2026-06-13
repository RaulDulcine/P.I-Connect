from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL   = RGBColor(0x1E, 0x40, 0xAF)
CINZA  = RGBColor(0x6B, 0x72, 0x80)
PRETO  = RGBColor(0x11, 0x18, 0x27)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
CABEC  = RGBColor(0x1E, 0x40, 0xAF)
ALT    = RGBColor(0xF9, 0xFA, 0xFB)



def bg(cell, r, g, b):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    s  = OxmlElement('w:shd')
    s.set(qn('w:val'),   'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'),  f'{r:02X}{g:02X}{b:02X}')
    pr.append(s)


def h(doc, text, level=1):
    para = doc.add_paragraph()
    sizes  = {1: 15, 2: 12, 3: 11}
    spaces = {1: (14, 6), 2: (10, 4), 3: (8, 3)}
    sb, sa = spaces.get(level, (8, 3))
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    run = para.add_run(text)
    run.font.bold      = True
    run.font.size      = Pt(sizes.get(level, 11))
    run.font.color.rgb = AZUL
    return para


def p(doc, text='', bold=False, italic=False, color=None, size=10, sa=6, sb=2, indent=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(sa)
    para.paragraph_format.space_before = Pt(sb)
    if indent:
        para.paragraph_format.left_indent = Cm(indent)
    if text:
        run = para.add_run(text)
        run.font.size   = Pt(size)
        run.font.bold   = bold
        run.font.italic = italic
        run.font.color.rgb = color or PRETO
    return para


def bullet(doc, text, size=10):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.space_before = Pt(1)
    run = para.add_run(text)
    run.font.size = Pt(size)
    return para


def tbl(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Cabeçalho
    hr = table.rows[0]
    for i, h_text in enumerate(headers):
        cell = hr.cells[i]
        bg(cell, 0x1E, 0x40, 0xAF)
        run = cell.paragraphs[0].add_run(h_text)
        run.font.bold      = True
        run.font.size      = Pt(9)
        run.font.color.rgb = BRANCO
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    # Dados
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        if ri % 2 == 1:
            for cell in row.cells:
                bg(cell, 0xF9, 0xFA, 0xFB)
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            run  = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def code(doc, text):
    for line in text.strip('\n').split('\n'):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent  = Cm(0.4)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run(line if line.strip() else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = PRETO
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),  'clear')
        shd.set(qn('w:fill'), 'F4F4F5')
        run._r.get_or_add_rPr().append(shd)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)


para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('PI Connect')
run.font.size = Pt(28); run.font.bold = True; run.font.color.rgb = AZUL

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('Relatório de Modelagem de Banco de Dados')
run.font.size = Pt(16); run.font.color.rgb = CINZA

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('Plataforma de Gestão de Projetos Integradores')
run.font.size = Pt(12); run.font.color.rgb = CINZA

doc.add_paragraph()

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('Versão 1.0  |   Junho de 2026')
run.font.size = Pt(10); run.font.color.rgb = CINZA

doc.add_page_break()


h(doc, '1. Introdução')
h(doc, '1.1 Objetivo do Sistema', 2)
p(doc, ('O PI Connect é uma plataforma web de gestão de Projetos Integradores (PI) desenvolvida para '
        'ambientes acadêmicos. O sistema centraliza o ciclo de vida dos projetos estudantis, desde a '
        'submissão pelo aluno até a aprovação final pelo professor orientador, com visibilidade gerencial '
        'para o administrador institucional.'))

h(doc, '1.2 Visão Geral da Solução', 2)
p(doc, 'A plataforma opera com três perfis de usuário distintos, cada um com painel dedicado:')
bullet(doc, 'Administrador – visão macro da plataforma (usuários, projetos, taxas de aprovação, projetos por curso)')
bullet(doc, 'Professor – gerenciamento dos projetos sob sua orientação, avaliações e feed de atividade recente')
bullet(doc, 'Aluno (Estudante) – submissão, acompanhamento de status e portfólio de projetos pessoais')

h(doc, '1.3 Público-Alvo', 2)
p(doc, ('Instituições de ensino técnico e superior que utilizam Projetos Integradores como metodologia '
        'avaliativa, envolvendo alunos, professores orientadores e coordenadores/administradores de curso.'))

doc.add_page_break()


h(doc, '2. Levantamento de Requisitos')
h(doc, '2.1 Requisitos Funcionais', 2)
for rf in [
    'RF01 – Autenticação de usuários com controle de perfil (Admin, Professor, Aluno)',
    'RF02 – Cadastro e gerenciamento de usuários pelo Administrador',
    'RF03 – Gerenciamento completo de Projetos Integradores pelo Aluno: submissão (criação), visualização, edição/atualização de dados e arquivos, e exclusão de registros',
    'RF04 – Listagem e filtragem de projetos por status (Aprovado, Pendente, Rejeitado)',
    'RF05 – Avaliação de projetos pelo Professor orientador',
    'RF06 – Aprovação e solicitação de revisão de projetos com feedback detalhado',
    'RF07 – Visualização de dashboard personalizado por perfil de usuário',
    'RF08 – Exibição de projetos por curso (ADS, Jogos Digitais, etc.)',
    'RF09 – Feed de atividade recente no painel do Professor (envios, atualizações, comentários)',
    'RF10 – Portfólio público do Aluno com vitrine de projetos publicados, visível para empresas parceiras do Senac',
    'RF11 – Relatórios e estatísticas para o Administrador (evolução por ano, distribuição por tema, status por curso, tendência mensal)',
    'RF12 – Controle de projetos supervisionados por Professor com filtro por status e ano',
    'RF13 – Registro de comentários e histórico de feedback em projetos, com controle de visibilidade ao aluno',
    'RF14 – Edição de perfil do usuário: nome, telefone, localização, biografia, senha e setor/departamento (para Administrador)',
    'RF15 – Explorar projetos de outros alunos com busca por título, autor e tecnologia',
]:
    bullet(doc, rf)

h(doc, '2.2 Requisitos Não Funcionais', 2)
for rnf in [
    'RNF01 – Segurança: senhas armazenadas com hashing bcrypt (fator de custo ≥ 12)',
    'RNF02 – Segurança: autenticação via JWT com expiração configurável',
    'RNF03 – Controle de acesso baseado em papéis (RBAC – Role-Based Access Control)',
    'RNF04 – Conformidade com a LGPD: consentimento de uso de dados, anonimização e direito ao esquecimento',
    'RNF05 – Auditoria: registro de todas as ações relevantes em log de auditoria',
    'RNF06 – Desempenho: consultas de dashboard respondidas em menos de 500 ms',
    'RNF07 – Escalabilidade: modelo relacional normalizado para suportar múltiplos cursos e semestres',
    'RNF08 – Usabilidade: interface responsiva com indicadores visuais de status de projeto',
    'RNF09 – Disponibilidade: banco de dados com backup diário automatizado',
    'RNF10 – Integridade referencial: uso de chaves estrangeiras com ON DELETE controlado',
]:
    bullet(doc, rnf)

doc.add_page_break()


h(doc, '3. Entidades do Banco de Dados')

# 3.1 usuario
h(doc, '3.1 usuario', 2)
p(doc, ('Armazena todos os usuários do sistema independentemente do perfil. O campo tipo_perfil '
        'determina o papel do usuário e, consequentemente, as funcionalidades acessíveis. Campos de '
        'perfil público (telefone, localização, biografia) permitem a edição de perfil por todos os painéis. '
        'O campo setor é específico para o perfil Administrador, armazenando o setor ou departamento '
        'institucional ao qual pertence.'))
tbl(doc,
    ['Campo', 'Tipo', 'PK', 'FK', 'Null', 'Descrição'],
    [
        ['id_usuario',    'INT UNSIGNED',                        '✓', '',  'N', 'Identificador único auto-incremental'],
        ['nome',          'VARCHAR(150)',                         '',  '',  'N', 'Nome completo do usuário'],
        ['email',         'VARCHAR(255)',                         '',  '',  'N', 'E-mail institucional, único no sistema'],
        ['senha_hash',    'VARCHAR(255)',                         '',  '',  'N', 'Hash bcrypt da senha'],
        ['tipo_perfil',   "ENUM('admin','professor','aluno')",    '',  '',  'N', 'Papel do usuário no sistema'],
        ['telefone',      'VARCHAR(20)',                          '',  '',  'S', 'Telefone de contato'],
        ['localizacao',   'VARCHAR(100)',                         '',  '',  'S', 'Cidade e estado (ex: Recife, PE)'],
        ['biografia',     'TEXT',                                 '',  '',  'S', 'Texto de apresentação do perfil público'],
        ['foto_url',      'VARCHAR(500)',                         '',  '',  'S', 'URL da foto de perfil'],
        ['setor',         'VARCHAR(100)',                         '',  '',  'S', 'Setor ou departamento institucional (usado no perfil do Administrador)'],
        ['status',        'TINYINT(1)',                           '',  '',  'N', '1 = ativo, 0 = inativo'],
        ['data_cadastro', 'DATETIME',                             '',  '',  'N', 'Data/hora de criação da conta'],
        ['ultimo_acesso', 'DATETIME',                             '',  '',  'S', 'Último login registrado'],
    ],
    [1.6, 2.0, 0.4, 0.4, 0.5, 2.5])

# 3.2 curso
h(doc, '3.2 curso', 2)
p(doc, ('Representa os cursos da instituição. Observado nas telas: ADS (Análise e Desenvolvimento de '
        'Sistemas) e Jogos Digitais. Modelado separadamente para permitir expansão sem alteração de esquema.'))
tbl(doc,
    ['Campo', 'Tipo', 'PK', 'FK', 'Null', 'Descrição'],
    [
        ['id_curso', 'INT UNSIGNED',  '✓', '', 'N', 'Identificador único'],
        ['nome',     'VARCHAR(150)',  '',  '', 'N', 'Nome do curso (ex: Análise e Desenvolvimento de Sistemas)'],
        ['sigla',    'VARCHAR(20)',   '',  '', 'N', 'Sigla do curso (ex: ADS)'],
        ['ativo',    'TINYINT(1)',    '',  '', 'N', '1 = curso ativo na plataforma'],
    ],
    [1.4, 1.6, 0.4, 0.4, 0.5, 3.1])

# 3.3 aluno
h(doc, '3.3 aluno', 2)
p(doc, ('Estende o usuário com dados específicos do perfil estudantil. Criado automaticamente via '
        'signal quando um usuário com tipo_perfil="aluno" é cadastrado, permitindo que o perfil '
        'seja completado posteriormente.'))
tbl(doc,
    ['Campo', 'Tipo', 'PK', 'FK', 'Null', 'Descrição'],
    [
        ['id_aluno',           'INT UNSIGNED', '✓', '',  'N', 'Identificador único'],
        ['id_usuario',         'INT UNSIGNED', '',  '✓', 'N', 'FK → usuario.id_usuario'],
        ['id_curso',           'INT UNSIGNED', '',  '✓', 'S', 'FK → curso.id_curso'],
        ['matricula',          'VARCHAR(20)',   '',  '',  'S', 'Número de matrícula institucional'],
        ['semestre_ingresso',  'YEAR',          '',  '',  'S', 'Ano de ingresso no curso'],
    ],
    [1.4, 1.6, 0.4, 0.4, 0.5, 3.1])

# 3.4 professor
h(doc, '3.4 professor', 2)
p(doc, ('Estende o usuário com dados do perfil docente. Professores supervisionam projetos e realizam '
        'avaliações. O campo id_curso indica o departamento de atuação principal, exibido na '
        'Gestão de Usuários do painel administrativo.'))
tbl(doc,
    ['Campo', 'Tipo', 'PK', 'FK', 'Null', 'Descrição'],
    [
        ['id_professor',        'INT UNSIGNED',  '✓', '',  'N', 'Identificador único'],
        ['id_usuario',          'INT UNSIGNED',  '',  '✓', 'N', 'FK → usuario.id_usuario'],
        ['id_curso',            'INT UNSIGNED',  '',  '✓', 'S', 'FK → curso.id_curso (departamento)'],
        ['titulacao',           'VARCHAR(100)',   '',  '',  'S', 'Titulação acadêmica (ex: Mestre, Doutor)'],
        ['registro_funcional',  'VARCHAR(30)',    '',  '',  'S', 'Registro funcional institucional'],
    ],
    [1.6, 1.6, 0.4, 0.4, 0.5, 2.9])

doc.add_page_break()

# 3.5 projeto
h(doc, '3.5 projeto', 2)
p(doc, ('Entidade central do sistema. Armazena todos os projetos integradores submetidos pelos alunos. '
        'Campos de classificação (tema, tecnologias, ano) alimentam os filtros da tela Explorar e os '
        'gráficos de relatório. O campo status_projeto controla o fluxo de aprovação e é central para '
        'os dashboards de todos os perfis.'))
tbl(doc,
    ['Campo', 'Tipo', 'PK', 'FK', 'Null', 'Descrição'],
    [
        ['id_projeto',        'INT UNSIGNED',    '✓', '',  'N', 'Identificador único'],
        ['titulo',            'VARCHAR(255)',     '',  '',  'N', 'Título do projeto'],
        ['descricao',         'TEXT',             '',  '',  'S', 'Descrição detalhada do projeto'],
        ['objetivos',         'TEXT',             '',  '',  'S', 'Objetivos do projeto (um por linha)'],
        ['tema',              'VARCHAR(10)',       '',  '',  'N', 'Classificação ESG/ODS (ex: ODS_4, ESG_AMB)'],
        ['tecnologias',       'VARCHAR(500)',      '',  '',  'S', 'Tecnologias utilizadas, separadas por vírgula'],
        ['ano',               'SMALLINT UNSIGNED','',  '',  'N', 'Ano de desenvolvimento do projeto'],
        ['id_aluno',          'INT UNSIGNED',     '',  '✓', 'N', 'FK → aluno.id_aluno (autor)'],
        ['id_professor',      'INT UNSIGNED',     '',  '✓', 'S', 'FK → professor.id_professor (orientador)'],
        ['id_curso',          'INT UNSIGNED',     '',  '✓', 'N', 'FK → curso.id_curso'],
        ['status_projeto',    'ENUM',             '',  '',  'N', 'pendente | aguardando_avaliacao | aprovado | rejeitado'],
        ['arquivo',           'VARCHAR(500)',      '',  '',  'S', 'Path do arquivo entregável (PDF, ZIP)'],
        ['visivel_portfolio',  'TINYINT(1)',       '',  '',  'N', '1 = exibido no portfólio público e em Explorar'],
        ['visualizacoes',     'INT UNSIGNED',      '',  '',  'N', 'Contador de acessos ao projeto em Explorar'],
        ['data_submissao',    'DATETIME',          '',  '',  'N', 'Data/hora de envio pelo aluno'],
        ['data_atualizacao',  'DATETIME',          '',  '',  'S', 'Última modificação do projeto'],
        ['data_avaliacao',    'DATETIME',          '',  '',  'S', 'Data/hora da avaliação pelo professor'],
    ],
    [1.6, 1.6, 0.4, 0.4, 0.5, 2.9])

# 3.6 avaliacao
h(doc, '3.6 avaliacao', 2)
p(doc, ('Registra o parecer formal do professor sobre um projeto. Separada da entidade projeto para '
        'preservar histórico de múltiplas avaliações (ex: resubmissão após revisão solicitada). '
        'Ao ser salva, propaga automaticamente a decisão para o status do projeto.'))
tbl(doc,
    ['Campo', 'Tipo', 'PK', 'FK', 'Null', 'Descrição'],
    [
        ['id_avaliacao',   'INT UNSIGNED',                      '✓', '',  'N', 'Identificador único'],
        ['id_projeto',     'INT UNSIGNED',                      '',  '✓', 'N', 'FK → projeto.id_projeto'],
        ['id_professor',   'INT UNSIGNED',                      '',  '✓', 'N', 'FK → professor.id_professor'],
        ['decisao',        "ENUM('aprovado','solicitar_revisao')",'', '', 'N', 'Resultado da avaliação'],
        ['feedback',       'TEXT',                               '',  '',  'N', 'Comentário/justificativa visível ao aluno'],
        ['data_avaliacao', 'DATETIME',                           '',  '',  'N', 'Data/hora do registro'],
    ],
    [1.4, 2.2, 0.4, 0.4, 0.5, 2.5])

doc.add_page_break()

# 3.7 comentario
h(doc, '3.7 comentario', 2)
p(doc, ('Registra interações textuais entre alunos e professores no contexto de um projeto. '
        'Evidenciado na seção "Histórico de Comentários" da tela de Avaliações e no feed '
        '"Atividade Recente" do dashboard do professor. Todos os perfis autenticados podem comentar.'))
tbl(doc,
    ['Campo', 'Tipo', 'PK', 'FK', 'Null', 'Descrição'],
    [
        ['id_comentario',  'INT UNSIGNED', '✓', '',  'N', 'Identificador único'],
        ['id_projeto',     'INT UNSIGNED', '',  '✓', 'N', 'FK → projeto.id_projeto'],
        ['id_usuario',     'INT UNSIGNED', '',  '✓', 'N', 'FK → usuario.id_usuario (autor)'],
        ['conteudo',       'TEXT',          '',  '',  'N', 'Conteúdo do comentário'],
        ['data_criacao',   'DATETIME',      '',  '',  'N', 'Data/hora de criação'],
        ['visivel_aluno',  'TINYINT(1)',    '',  '',  'N', '1 = visível para o aluno (padrão)'],
    ],
    [1.4, 1.6, 0.4, 0.4, 0.5, 3.1])

# 3.8 atividade_log
h(doc, '3.8 atividade_log', 2)
p(doc, ('Registra eventos relevantes para o feed "Atividade Recente" e auditoria geral. '
        'tipo_acao categoriza o evento (envio, atualização, comentário, avaliação).'))
tbl(doc,
    ['Campo', 'Tipo', 'PK', 'FK', 'Null'],
    [
        ['id_log',     'BIGINT UNSIGNED', '✓', '',  'N'],
        ['id_usuario', 'INT UNSIGNED',    '',  '✓', 'N'],
        ['id_projeto', 'INT UNSIGNED',    '',  '✓', 'S'],
        ['tipo_acao',  "ENUM('envio','atualizacao','comentario','avaliacao','login','cadastro')", '', '', 'N'],
        ['descricao',  'VARCHAR(500)',     '',  '',  'S'],
        ['data_acao',  'DATETIME',         '',  '',  'N'],
        ['ip_origem',  'VARCHAR(45)',      '',  '',  'S'],
    ],
    [1.4, 2.8, 0.4, 0.4, 0.5])

# 3.9 sessao
h(doc, '3.9 sessao', 2)
p(doc, ('Gerencia tokens de sessão/JWT para controle de autenticação ativa. '
        'Permite invalidação de sessões (logout em todos os dispositivos) e auditoria de acesso.'))
tbl(doc,
    ['Campo', 'Tipo', 'PK', 'FK', 'Null', 'Descrição'],
    [
        ['id_sessao',   'VARCHAR(64)',  '✓', '',  'N', 'Token JWT ou UUID de sessão'],
        ['id_usuario',  'INT UNSIGNED', '',  '✓', 'N', 'FK → usuario.id_usuario'],
        ['criado_em',   'DATETIME',     '',  '',  'N', 'Data/hora de criação da sessão'],
        ['expira_em',   'DATETIME',     '',  '',  'N', 'Data/hora de expiração'],
        ['invalidado',  'TINYINT(1)',   '',  '',  'N', '1 = sessão invalidada (logout)'],
        ['user_agent',  'VARCHAR(500)', '',  '',  'S', 'Navegador/dispositivo do cliente'],
    ],
    [1.4, 1.6, 0.4, 0.4, 0.5, 3.1])

doc.add_page_break()


h(doc, '4. Modelo Conceitual')
p(doc, 'O modelo conceitual identifica as entidades principais e seus relacionamentos, '
       'expressando as regras de negócio da plataforma PI Connect:')
tbl(doc,
    ['Relacionamento', 'Cardinalidade', 'Tipo', 'Regra de Negócio'],
    [
        ['usuario → aluno / professor', '1 : 0..1', 'Especialização',    'Um usuário possui exatamente um perfil especializado (aluno ou professor)'],
        ['aluno → curso',               'N : 1',    'Associação',        'Cada aluno pertence a um curso; um curso tem muitos alunos'],
        ['professor → curso',           'N : 1',    'Associação',        'Cada professor está vinculado a um departamento/curso'],
        ['aluno → projeto',             '1 : N',    'Composição',        'Um aluno pode submeter múltiplos projetos; cada projeto tem exatamente um autor'],
        ['professor → projeto',         '1 : N',    'Associação',        'Um professor orienta vários projetos; um projeto tem no máximo um orientador'],
        ['projeto → avaliacao',         '1 : N',    'Associação',        'Um projeto pode ter múltiplas avaliações (resubmissões)'],
        ['professor → avaliacao',       '1 : N',    'Associação',        'Um professor realiza muitas avaliações'],
        ['projeto → comentario',        '1 : N',    'Agregação',         'Um projeto acumula vários comentários de professores e alunos'],
        ['usuario → comentario',        '1 : N',    'Associação',        'Qualquer usuário autenticado pode comentar'],
        ['usuario → atividade_log',     '1 : N',    'Associação',        'Todo evento registrado aponta para o usuário que o gerou'],
        ['projeto → atividade_log',     '1 : N',    'Associação',        'Eventos contextualizados em projetos referenciam o projeto'],
        ['curso → projeto',             '1 : N',    'Associação',        'Projetos são vinculados ao curso do aluno para gráficos por curso'],
    ],
    [2.0, 1.2, 1.2, 3.0])

doc.add_page_break()


h(doc, '5. Modelo Lógico')
p(doc, 'O modelo lógico normaliza as entidades para implementação relacional em MySQL 8.x, '
       'respeitando integridade referencial e os padrões observados nas telas do protótipo.')
tbl(doc,
    ['Tabela', 'PK', 'FKs', 'Restrições / Índices'],
    [
        ['usuario',       'id_usuario (PK)',      '—',                                               'UNIQUE(email), INDEX(tipo_perfil)'],
        ['curso',         'id_curso (PK)',         '—',                                               'UNIQUE(sigla)'],
        ['aluno',         'id_aluno (PK)',          'id_usuario (FK), id_curso (FK)',                  'UNIQUE(id_usuario), UNIQUE(matricula)'],
        ['professor',     'id_professor (PK)',      'id_usuario (FK), id_curso (FK)',                  'UNIQUE(id_usuario)'],
        ['projeto',       'id_projeto (PK)',        'id_aluno (FK), id_professor (FK), id_curso (FK)', 'INDEX(status_projeto), INDEX(data_submissao), INDEX(id_curso), INDEX(visivel_portfolio, visualizacoes)'],
        ['avaliacao',     'id_avaliacao (PK)',      'id_projeto (FK), id_professor (FK)',              'INDEX(id_projeto), INDEX(data_avaliacao)'],
        ['comentario',    'id_comentario (PK)',     'id_projeto (FK), id_usuario (FK)',                'INDEX(id_projeto), INDEX(data_criacao)'],
        ['atividade_log', 'id_log (PK)',            'id_usuario (FK), id_projeto (FK)',                'INDEX(tipo_acao, data_acao), INDEX(id_projeto)'],
        ['sessao',        'id_sessao (PK)',         'id_usuario (FK)',                                  'INDEX(id_usuario, expira_em)'],
    ],
    [1.3, 1.5, 2.2, 2.4])

doc.add_page_break()


h(doc, '6. Dicionário de Dados')
p(doc, 'Campos selecionados de maior relevância para o negócio, agrupados por tabela.')

p(doc, 'Tabela: projeto', bold=True)
tbl(doc,
    ['Campo', 'Tipo', 'Obrigatório', 'Descrição'],
    [
        ['status_projeto',    'ENUM',         'Sim', 'Controla o fluxo de aprovação. Valores: pendente (submetido, aguardando atribuição), aguardando_avaliacao (professor atribuído), aprovado, rejeitado.'],
        ['tema',              'VARCHAR(10)',   'Sim', 'Classificação ESG/ODS do projeto. Valores: ODS_4, ODS_7, ODS_10, ODS_11, ODS_12, ODS_13, ESG_AMB, ESG_SOC, ESG_GOV.'],
        ['data_submissao',    'DATETIME',      'Sim', 'Timestamp gerado automaticamente no INSERT. Não editável pelo usuário.'],
        ['visivel_portfolio', 'TINYINT(1)',    'Sim', 'Flag que determina se o projeto aparece no portfólio público do aluno e na tela Explorar. Ativado automaticamente ao aprovar.'],
        ['visualizacoes',     'INT UNSIGNED',  'Sim', 'Incrementado a cada acesso à página do projeto. Exibido em Explorar e Portfólio.'],
        ['arquivo',           'VARCHAR(500)',  'Não', 'Path relativo do arquivo entregável (PDF, ZIP) armazenado no servidor.'],
        ['objetivos',         'TEXT',          'Não', 'Lista de objetivos do projeto, um por linha. Exibida na tela de Avaliações.'],
    ],
    [1.5, 1.3, 1.0, 3.6])

p(doc, 'Tabela: usuario', bold=True)
tbl(doc,
    ['Campo', 'Tipo', 'Obrigatório', 'Descrição'],
    [
        ['senha_hash',   'VARCHAR(255)', 'Sim', 'Nunca armazenada em texto plano. Hash bcrypt com fator de custo mínimo 12. Nunca retornada em consultas de API.'],
        ['tipo_perfil',  'ENUM',         'Sim', 'Discriminador de perfil: admin, professor, aluno. Define quais módulos o usuário acessa.'],
        ['localizacao',  'VARCHAR(100)', 'Não', 'Cidade e estado informados pelo usuário na edição de perfil. Ex: Recife, PE.'],
        ['biografia',    'TEXT',         'Não', 'Texto de apresentação exibido no perfil público do usuário.'],
        ['setor',        'VARCHAR(100)', 'Não', 'Setor ou departamento institucional. Exibido e editável apenas no perfil do Administrador. Ex: Coordenação Acadêmica.'],
    ],
    [1.5, 1.3, 1.0, 3.6])

p(doc, 'Tabela: avaliacao', bold=True)
tbl(doc,
    ['Campo', 'Tipo', 'Obrigatório', 'Descrição'],
    [
        ['decisao',  'ENUM',  'Sim', 'Resultado: aprovado (projeto entra no portfólio) ou solicitar_revisao (aluno deve revisar e resubmeter).'],
        ['feedback', 'TEXT',  'Sim', 'Comentário do professor visível ao aluno. Mínimo de 20 caracteres validado pela aplicação.'],
    ],
    [1.5, 1.3, 1.0, 3.6])

p(doc, 'Tabela: atividade_log', bold=True)
tbl(doc,
    ['Campo', 'Tipo', 'Obrigatório', 'Descrição'],
    [
        ['tipo_acao', 'ENUM',        'Sim', 'Categoriza o evento para filtragem. O dashboard do professor exibe apenas ações relacionadas a seus projetos.'],
        ['descricao', 'VARCHAR(500)','Não', 'Texto legível humanamente gerado pela aplicação. Ex: "João Vitor enviou Sistema de Gestão Escolar".'],
        ['ip_origem', 'VARCHAR(45)', 'Não', 'Registrado para fins de auditoria de segurança. IPv6 requer até 45 caracteres.'],
    ],
    [1.5, 1.3, 1.0, 3.6])

doc.add_page_break()


h(doc, '7. Normalização')

h(doc, '7.1 Primeira Forma Normal (1FN)', 2)
p(doc, ('Todas as tabelas possuem chave primária definida e todos os atributos são atômicos (valores '
        'indivisíveis). Não há grupos repetitivos: os cursos de um aluno não são armazenados em campos '
        'separados (id_curso1, id_curso2), mas em uma tabela específica com FK. O ENUM status_projeto '
        'representa um conjunto fechado de valores atômicos, não uma lista. O campo tecnologias '
        'armazena uma lista como texto livre (CSV), tratado como atômico na camada de aplicação; '
        'a normalização completa deste campo está prevista nas Melhorias Futuras (tabela tag).'))

h(doc, '7.2 Segunda Forma Normal (2FN)', 2)
p(doc, ('O modelo não utiliza chaves compostas nas tabelas principais; todas as PKs são surrogate keys '
        '(INT UNSIGNED AUTO_INCREMENT). Portanto, não há dependências parciais possíveis. Todos os '
        'atributos dependem funcionalmente da chave primária inteira. Exemplo: o campo titulo em projeto '
        'depende exclusivamente de id_projeto, não de uma combinação parcial.'))

h(doc, '7.3 Terceira Forma Normal (3FN)', 2)
p(doc, ('Não há dependências transitivas. Exemplo de decisão de modelagem que garante 3FN: o nome do '
        'curso não é armazenado em projeto (o que criaria dependência transitiva id_projeto → id_curso → '
        'nome_curso). Em vez disso, projeto.id_curso referencia a tabela curso, onde o nome é armazenado '
        'uma única vez. O mesmo princípio aplica-se a nome do professor e nome do aluno, que residem em '
        'usuario e são acessados via JOIN.'))

doc.add_page_break()


h(doc, '8. Diagrama Entidade-Relacionamento (DER)')
p(doc, 'O diagrama abaixo, em notação Mermaid (ERD), representa as entidades e relacionamentos identificados:')

code(doc, """
erDiagram
    usuario {
        INT       id_usuario    PK
        VARCHAR   nome
        VARCHAR   email
        VARCHAR   senha_hash
        ENUM      tipo_perfil
        VARCHAR   telefone
        VARCHAR   localizacao
        TEXT      biografia
        VARCHAR   foto_url
        VARCHAR   setor
        TINYINT   status
        DATETIME  data_cadastro
    }
    curso {
        INT       id_curso  PK
        VARCHAR   nome
        VARCHAR   sigla
        TINYINT   ativo
    }
    aluno {
        INT       id_aluno           PK
        INT       id_usuario         FK
        INT       id_curso           FK
        VARCHAR   matricula
        YEAR      semestre_ingresso
    }
    professor {
        INT       id_professor        PK
        INT       id_usuario          FK
        INT       id_curso            FK
        VARCHAR   titulacao
        VARCHAR   registro_funcional
    }
    projeto {
        INT       id_projeto          PK
        VARCHAR   titulo
        TEXT      descricao
        TEXT      objetivos
        VARCHAR   tema
        VARCHAR   tecnologias
        SMALLINT  ano
        INT       id_aluno            FK
        INT       id_professor        FK
        INT       id_curso            FK
        ENUM      status_projeto
        DATETIME  data_submissao
        DATETIME  data_atualizacao
        VARCHAR   arquivo
        TINYINT   visivel_portfolio
        INT       visualizacoes
    }
    avaliacao {
        INT       id_avaliacao    PK
        INT       id_projeto      FK
        INT       id_professor    FK
        ENUM      decisao
        TEXT      feedback
        DATETIME  data_avaliacao
    }
    comentario {
        INT       id_comentario  PK
        INT       id_projeto     FK
        INT       id_usuario     FK
        TEXT      conteudo
        DATETIME  data_criacao
        TINYINT   visivel_aluno
    }
    atividade_log {
        BIGINT    id_log       PK
        INT       id_usuario   FK
        INT       id_projeto   FK
        ENUM      tipo_acao
        VARCHAR   descricao
        DATETIME  data_acao
        VARCHAR   ip_origem
    }
    sessao {
        VARCHAR   id_sessao   PK
        INT       id_usuario  FK
        DATETIME  criado_em
        DATETIME  expira_em
        TINYINT   invalidado
    }

    usuario      ||--o|  aluno         : "especializa"
    usuario      ||--o|  professor     : "especializa"
    aluno        }o--||  curso         : "pertence a"
    professor    }o--o|  curso         : "atua em"
    aluno        ||--o{  projeto       : "submete"
    professor    ||--o{  projeto       : "orienta"
    curso        ||--o{  projeto       : "categoriza"
    projeto      ||--o{  avaliacao     : "recebe"
    professor    ||--o{  avaliacao     : "realiza"
    projeto      ||--o{  comentario    : "possui"
    usuario      ||--o{  comentario    : "escreve"
    usuario      ||--o{  atividade_log : "gera"
    projeto      ||--o{  atividade_log : "contextualiza"
    usuario      ||--o{  sessao        : "mantém"
""")

p(doc, 'Renderize o código acima em https://mermaid.live para visualizar o diagrama interativo.',
  italic=True, color=CINZA)

doc.add_page_break()


h(doc, '9. Scripts SQL')
p(doc, 'Scripts compatíveis com MySQL 8.x. Executar na ordem apresentada para respeitar dependências de FK.')

code(doc, """
-- ============================================================
-- PI Connect – Script de Criação do Banco de Dados
-- MySQL 8.x | ENGINE InnoDB | Charset utf8mb4
-- ============================================================

CREATE DATABASE IF NOT EXISTS pi_connect
  CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;

USE pi_connect;

-- ------------------------------------------------------------
-- 1. curso
-- ------------------------------------------------------------
CREATE TABLE curso (
  id_curso  INT UNSIGNED NOT NULL AUTO_INCREMENT,
  nome      VARCHAR(150) NOT NULL,
  sigla     VARCHAR(20)  NOT NULL,
  ativo     TINYINT(1)   NOT NULL DEFAULT 1,
  PRIMARY KEY (id_curso),
  UNIQUE KEY uq_curso_sigla (sigla)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;

-- ------------------------------------------------------------
-- 2. usuario
-- ------------------------------------------------------------
CREATE TABLE usuario (
  id_usuario    INT UNSIGNED  NOT NULL AUTO_INCREMENT,
  nome          VARCHAR(150)  NOT NULL,
  email         VARCHAR(255)  NOT NULL,
  senha_hash    VARCHAR(255)  NOT NULL,
  tipo_perfil   ENUM('admin','professor','aluno') NOT NULL,
  telefone      VARCHAR(20)   NULL,
  localizacao   VARCHAR(100)  NULL,
  biografia     TEXT          NULL,
  foto_url      VARCHAR(500)  NULL,
  setor         VARCHAR(100)  NULL,
  status        TINYINT(1)    NOT NULL DEFAULT 1,
  data_cadastro DATETIME      NOT NULL DEFAULT CURRENT_TIMESTAMP,
  ultimo_acesso DATETIME      NULL,
  PRIMARY KEY (id_usuario),
  UNIQUE KEY uq_usuario_email (email),
  INDEX idx_usuario_tipo (tipo_perfil)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;

-- ------------------------------------------------------------
-- 3. aluno
-- ------------------------------------------------------------
CREATE TABLE aluno (
  id_aluno          INT UNSIGNED NOT NULL AUTO_INCREMENT,
  id_usuario        INT UNSIGNED NOT NULL,
  id_curso          INT UNSIGNED NULL,
  matricula         VARCHAR(20)  NULL,
  semestre_ingresso YEAR         NULL,
  PRIMARY KEY (id_aluno),
  UNIQUE KEY uq_aluno_usuario  (id_usuario),
  UNIQUE KEY uq_aluno_matricula (matricula),
  CONSTRAINT fk_aluno_usuario FOREIGN KEY (id_usuario)
    REFERENCES usuario (id_usuario) ON DELETE CASCADE   ON UPDATE CASCADE,
  CONSTRAINT fk_aluno_curso   FOREIGN KEY (id_curso)
    REFERENCES curso   (id_curso)   ON DELETE SET NULL  ON UPDATE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;

-- ------------------------------------------------------------
-- 4. professor
-- ------------------------------------------------------------
CREATE TABLE professor (
  id_professor       INT UNSIGNED NOT NULL AUTO_INCREMENT,
  id_usuario         INT UNSIGNED NOT NULL,
  id_curso           INT UNSIGNED NULL,
  titulacao          VARCHAR(100) NULL,
  registro_funcional VARCHAR(30)  NULL,
  PRIMARY KEY (id_professor),
  UNIQUE KEY uq_professor_usuario (id_usuario),
  CONSTRAINT fk_professor_usuario FOREIGN KEY (id_usuario)
    REFERENCES usuario (id_usuario) ON DELETE CASCADE  ON UPDATE CASCADE,
  CONSTRAINT fk_professor_curso   FOREIGN KEY (id_curso)
    REFERENCES curso   (id_curso)   ON DELETE SET NULL ON UPDATE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;

-- ------------------------------------------------------------
-- 5. projeto
-- ------------------------------------------------------------
CREATE TABLE projeto (
  id_projeto        INT UNSIGNED    NOT NULL AUTO_INCREMENT,
  titulo            VARCHAR(255)    NOT NULL,
  descricao         TEXT            NULL,
  objetivos         TEXT            NULL,
  tema              VARCHAR(10)     NOT NULL,
  tecnologias       VARCHAR(500)    NULL,
  ano               SMALLINT UNSIGNED NOT NULL,
  id_aluno          INT UNSIGNED    NOT NULL,
  id_professor      INT UNSIGNED    NULL,
  id_curso          INT UNSIGNED    NOT NULL,
  status_projeto    ENUM('pendente','aguardando_avaliacao','aprovado','rejeitado')
                                    NOT NULL DEFAULT 'pendente',
  arquivo           VARCHAR(500)    NULL,
  visivel_portfolio TINYINT(1)      NOT NULL DEFAULT 0,
  visualizacoes     INT UNSIGNED    NOT NULL DEFAULT 0,
  data_submissao    DATETIME        NOT NULL DEFAULT CURRENT_TIMESTAMP,
  data_atualizacao  DATETIME        NULL ON UPDATE CURRENT_TIMESTAMP,
  data_avaliacao    DATETIME        NULL,
  PRIMARY KEY (id_projeto),
  INDEX idx_projeto_status    (status_projeto),
  INDEX idx_projeto_data      (data_submissao),
  INDEX idx_projeto_curso     (id_curso),
  INDEX idx_projeto_explorar  (visivel_portfolio, visualizacoes),
  CONSTRAINT fk_projeto_aluno     FOREIGN KEY (id_aluno)
    REFERENCES aluno     (id_aluno)     ON DELETE RESTRICT ON UPDATE CASCADE,
  CONSTRAINT fk_projeto_professor FOREIGN KEY (id_professor)
    REFERENCES professor (id_professor) ON DELETE SET NULL ON UPDATE CASCADE,
  CONSTRAINT fk_projeto_curso     FOREIGN KEY (id_curso)
    REFERENCES curso     (id_curso)     ON DELETE RESTRICT ON UPDATE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;

-- ------------------------------------------------------------
-- 6. avaliacao
-- ------------------------------------------------------------
CREATE TABLE avaliacao (
  id_avaliacao  INT UNSIGNED NOT NULL AUTO_INCREMENT,
  id_projeto    INT UNSIGNED NOT NULL,
  id_professor  INT UNSIGNED NOT NULL,
  decisao       ENUM('aprovado','solicitar_revisao') NOT NULL,
  feedback      TEXT         NOT NULL,
  data_avaliacao DATETIME    NOT NULL DEFAULT CURRENT_TIMESTAMP,
  PRIMARY KEY (id_avaliacao),
  INDEX idx_aval_projeto (id_projeto),
  INDEX idx_aval_data    (data_avaliacao),
  CONSTRAINT fk_aval_projeto   FOREIGN KEY (id_projeto)
    REFERENCES projeto   (id_projeto)   ON DELETE CASCADE  ON UPDATE CASCADE,
  CONSTRAINT fk_aval_professor FOREIGN KEY (id_professor)
    REFERENCES professor (id_professor) ON DELETE RESTRICT ON UPDATE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;

-- ------------------------------------------------------------
-- 7. comentario
-- ------------------------------------------------------------
CREATE TABLE comentario (
  id_comentario INT UNSIGNED NOT NULL AUTO_INCREMENT,
  id_projeto    INT UNSIGNED NOT NULL,
  id_usuario    INT UNSIGNED NOT NULL,
  conteudo      TEXT         NOT NULL,
  data_criacao  DATETIME     NOT NULL DEFAULT CURRENT_TIMESTAMP,
  visivel_aluno TINYINT(1)   NOT NULL DEFAULT 1,
  PRIMARY KEY (id_comentario),
  INDEX idx_coment_projeto (id_projeto),
  INDEX idx_coment_data    (data_criacao),
  CONSTRAINT fk_coment_projeto FOREIGN KEY (id_projeto)
    REFERENCES projeto  (id_projeto)  ON DELETE CASCADE  ON UPDATE CASCADE,
  CONSTRAINT fk_coment_usuario FOREIGN KEY (id_usuario)
    REFERENCES usuario  (id_usuario)  ON DELETE RESTRICT ON UPDATE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;

-- ------------------------------------------------------------
-- 8. atividade_log
-- ------------------------------------------------------------
CREATE TABLE atividade_log (
  id_log     BIGINT UNSIGNED NOT NULL AUTO_INCREMENT,
  id_usuario INT UNSIGNED    NOT NULL,
  id_projeto INT UNSIGNED    NULL,
  tipo_acao  ENUM('envio','atualizacao','comentario','avaliacao','login','cadastro')
                             NOT NULL,
  descricao  VARCHAR(500)    NULL,
  data_acao  DATETIME        NOT NULL DEFAULT CURRENT_TIMESTAMP,
  ip_origem  VARCHAR(45)     NULL,
  PRIMARY KEY (id_log),
  INDEX idx_log_usuario_data (id_usuario, data_acao),
  INDEX idx_log_projeto      (id_projeto),
  INDEX idx_log_tipo_data    (tipo_acao, data_acao),
  CONSTRAINT fk_log_usuario FOREIGN KEY (id_usuario)
    REFERENCES usuario  (id_usuario)  ON DELETE RESTRICT ON UPDATE CASCADE,
  CONSTRAINT fk_log_projeto FOREIGN KEY (id_projeto)
    REFERENCES projeto  (id_projeto)  ON DELETE SET NULL ON UPDATE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;

-- ------------------------------------------------------------
-- 9. sessao
-- ------------------------------------------------------------
CREATE TABLE sessao (
  id_sessao  VARCHAR(64)  NOT NULL,
  id_usuario INT UNSIGNED NOT NULL,
  criado_em  DATETIME     NOT NULL DEFAULT CURRENT_TIMESTAMP,
  expira_em  DATETIME     NOT NULL,
  invalidado TINYINT(1)   NOT NULL DEFAULT 0,
  user_agent VARCHAR(500) NULL,
  PRIMARY KEY (id_sessao),
  INDEX idx_sessao_usuario (id_usuario, expira_em),
  CONSTRAINT fk_sessao_usuario FOREIGN KEY (id_usuario)
    REFERENCES usuario (id_usuario) ON DELETE CASCADE ON UPDATE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;

-- ============================================================
-- Views úteis para dashboard
-- ============================================================

-- Contagem de projetos por status (Dashboard Administrativo)
CREATE OR REPLACE VIEW vw_projetos_por_status AS
SELECT status_projeto, COUNT(*) AS total
FROM projeto
GROUP BY status_projeto;

-- Projetos por curso (gráfico de barras – Dashboard Administrativo)
CREATE OR REPLACE VIEW vw_projetos_por_curso AS
SELECT c.sigla, c.nome, COUNT(p.id_projeto) AS total_projetos
FROM curso c
LEFT JOIN projeto p ON p.id_curso = c.id_curso
GROUP BY c.id_curso;

-- Taxa de aprovação global
CREATE OR REPLACE VIEW vw_taxa_aprovacao AS
SELECT
  ROUND(
    (COUNT(CASE WHEN status_projeto = 'aprovado' THEN 1 END) * 100.0)
    / NULLIF(COUNT(*), 0), 2
  ) AS taxa_aprovacao_pct
FROM projeto;

-- Evolução de projetos por ano (Relatórios)
CREATE OR REPLACE VIEW vw_projetos_por_ano AS
SELECT ano, COUNT(*) AS total
FROM projeto
GROUP BY ano
ORDER BY ano;

-- Distribuição por tema (Relatórios)
CREATE OR REPLACE VIEW vw_projetos_por_tema AS
SELECT tema, COUNT(*) AS total
FROM projeto
GROUP BY tema
ORDER BY total DESC;
""")

doc.add_page_break()


h(doc, '10. Considerações sobre Segurança')

h(doc, '10.1 Criptografia de Senhas', 2)
p(doc, ('Senhas são armazenadas exclusivamente como hashes bcrypt com fator de custo mínimo 12, '
        'tornando ataques de força bruta computacionalmente inviáveis. A senha nunca é retornada em '
        'nenhuma query de API ou listagem administrativa.'))

h(doc, '10.2 Controle de Acesso (RBAC)', 2)
p(doc, 'O campo tipo_perfil da tabela usuario determina o nível de acesso. A camada de aplicação valida o '
       'perfil em cada requisição por meio do token JWT. As regras são:')
bullet(doc, 'admin: acesso total a todos os módulos')
bullet(doc, 'professor: acesso apenas aos projetos vinculados via id_professor e às avaliações que realizou')
bullet(doc, 'aluno: acesso restrito aos próprios projetos e ao portfólio público')

h(doc, '10.3 Conformidade com a LGPD', 2)
p(doc, ('Os campos de dados pessoais (nome, email, telefone, localização) devem ter base legal registrada. '
        'Ao desativar um usuário (status = 0), a aplicação pode anonimizar nome e email substituindo por '
        'valores hash, preservando histórico estatístico sem identificação. O campo ip_origem em '
        'atividade_log requer consentimento explícito e política de retenção definida (sugerido: 90 dias).'))

h(doc, '10.4 Auditoria e Logs', 2)
p(doc, ('A tabela atividade_log registra todos os eventos relevantes: submissões, avaliações, comentários, '
        'logins e cadastros. O campo ip_origem e user_agent (em sessao) permitem rastrear acessos '
        'suspeitos. Recomenda-se manter a tabela atividade_log em partição por data para performance em '
        'alto volume.'))

h(doc, '10.5 Integridade Referencial', 2)
p(doc, ('Todas as FKs utilizam ENGINE InnoDB com ON DELETE controlado: RESTRICT para entidades que '
        'não devem ser excluídas enquanto possuírem dependentes (ex: projeto → aluno), CASCADE para '
        'dados filhos sem valor independente (ex: sessao → usuario), e SET NULL para referências '
        'opcionais (ex: projeto → professor, quando o orientador é removido).'))

doc.add_page_break()


h(doc, '11. Melhorias Futuras')
for mf in [
    'Links externos no perfil do Aluno: campos linkedin_url e github_url na tabela usuario, citados pelos usuários no levantamento de requisitos (10/11 votos) para exibição no portfólio público.',
    'Critérios de avaliação estruturados: tabela criterio_avaliacao com dimensões específicas (alinhamento com tema, viabilidade técnica, impacto social/ambiental, documentação, inovação), substituindo a decisão binária atual e alinhada aos critérios exibidos na tela de Avaliações.',
    'Tabela semestre: modelar semestres letivos (2026.1, 2026.2) e vincular projetos a semestres, permitindo relatórios históricos e comparativos entre períodos.',
    'Tabela tag / projeto_tag: normalização do campo tecnologias em uma entidade dedicada, melhorando filtros e relatórios por stack tecnológica na tela Explorar.',
    'Tabela notificacao: sistema de notificações persistentes vinculadas a eventos de atividade_log, substituindo o feed estático do dashboard.',
    'Tabela arquivo_versao: suporte a múltiplos arquivos por projeto (versões, documentação, apresentação), com versionamento e metadados (tamanho, tipo MIME).',
    'Tabela grupo_projeto: projetos desenvolvidos em equipe, com tabela associativa aluno_grupo para membros e papel (líder, colaborador).',
    'Tabela criterio_avaliacao / nota_criterio: rubricas de avaliação detalhadas por dimensão (inovação, qualidade técnica, apresentação), substituindo a decisão binária atual.',
    'Suporte a múltiplos perfis: um usuário poderia ser professor em um curso e aluno em outro (pós-graduação), requerendo refatoração da tabela aluno/professor para uma entidade perfil_usuario mais flexível.',
    'Particionamento da tabela atividade_log por RANGE em data_acao para performance em escala institucional.',
    'Integração com sistema de autenticação institucional (SSO/OAuth2), adicionando campo provider e provider_id à tabela usuario.',
]:
    bullet(doc, mf)

doc.add_paragraph()
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('Relatório gerado para a plataforma PI Connect  |   Junho de 2026')
run.font.size = Pt(9); run.font.color.rgb = CINZA


output = 'PI_Connect_Relatorio_Banco_de_Dados.docx'
doc.save(output)
print(f'Relatorio gerado: {output}')
